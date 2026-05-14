"""Geração do relatório semanal em .docx.

Recebe a mesma lista de processos (já redatada) que o .md usa e produz
um documento Word agrupado por órgão julgador, salvo junto do .md/.json.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REDATADO = "[REDATADO]"


def _fmt_data(iso: str | None) -> str:
    if not iso:
        return "—"
    try:
        return date.fromisoformat(str(iso)[:10]).strftime("%d/%m/%Y")
    except ValueError:
        return str(iso)


def _fmt_cnj(numero: str | None, fallback: str | None) -> str:
    s = "".join(ch for ch in (numero or "") if ch.isdigit())
    if len(s) == 20:
        return f"{s[:7]}-{s[7:9]}.{s[9:13]}.{s[13:14]}.{s[14:16]}.{s[16:20]}"
    return numero or fallback or "(sem número)"


def _fmt_partes(partes: dict | str | None) -> str:
    if partes is None:
        return "(sem partes informadas)"
    if isinstance(partes, str):  # já redatado
        return partes
    saidas: list[str] = []
    for polo in ("ativa", "passiva"):
        p = partes.get(polo)
        if not p:
            continue
        nome = p.get("nome_social") or p.get("nome") or "(sem nome)"
        tipo = p.get("tipo") or polo.capitalize()
        saidas.append(f"{tipo}: {nome}")
    return "  •  ".join(saidas) if saidas else "(sem partes informadas)"


def _campo(doc: Document, rotulo: str, valor: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{rotulo}: ").bold = True
    p.add_run(valor if valor not in (None, "") else "—")


def gerar_docx(
    processos: list[dict], *, de: date, ate: date, label: str, saida_path: Path
) -> Path:
    """Monta o .docx do relatório semanal, salva em saida_path e o retorna."""
    de_br = de.strftime("%d/%m/%Y")
    ate_br = ate.strftime("%d/%m/%Y")
    total = len(processos)
    redatados = sum(1 for p in processos if p["ementa"] == REDATADO)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("RELATÓRIO SEMANAL DE ACÓRDÃOS — TJMS")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Semana {label}  ·  {de_br} a {ate_br}").font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto_total = f"{total} acórdão(s) publicado(s)"
    if redatados:
        texto_total += f"  ({redatados} redatado(s) por privacidade)"
    meta.add_run(texto_total).font.size = Pt(11)

    doc.add_page_break()

    if total == 0:
        doc.add_paragraph("Nenhum acórdão publicado nesta semana.")
        saida_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(saida_path))
        return saida_path

    por_orgao: dict[int, list[dict]] = {}
    nomes: dict[int, str] = {}
    for p in processos:
        cd = p["orgao"]["cd"]
        por_orgao.setdefault(cd, []).append(p)
        nomes[cd] = p["orgao"]["nome"]

    for cd in sorted(por_orgao):
        grupo = por_orgao[cd]
        plural = "s" if len(grupo) > 1 else ""
        doc.add_heading(f"{nomes[cd]} ({len(grupo)} acórdão{plural})", level=1)
        for i, p in enumerate(grupo, 1):
            classe = p["classe"] or "Processo"
            numero = _fmt_cnj(p["numero_unificado"], p["codigo_processo"])
            tag = "  [REDATADO]" if p["ementa"] == REDATADO else ""
            doc.add_heading(f"{i}. {classe} nº {numero}{tag}", level=2)

            _campo(doc, "Relator(a)", p["relator"] or "(não informado)")
            _campo(doc, "Data de julgamento", _fmt_data(p["dt_julgamento"]))
            _campo(doc, "Publicação no DJE", _fmt_data(p["dt_publicacao_dje"]))
            _campo(doc, "Partes", _fmt_partes(p["partes"]))
            if p["segredo_justica"]:
                _campo(doc, "Segredo de justiça", "SIM")

            rot = doc.add_paragraph()
            rot.paragraph_format.space_before = Pt(6)
            rot.add_run("Ementa").bold = True

            ementa = p["ementa"] or "(ementa não disponível)"
            if ementa == REDATADO:
                doc.add_paragraph(REDATADO)
            else:
                for bloco in (b.strip() for b in ementa.replace("\r\n", "\n").split("\n")):
                    if bloco:
                        par = doc.add_paragraph(bloco)
                        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    saida_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida_path))
    return saida_path
